# report_builder.py — 使用python-docx生成Word报告

import sys
sys.stdout.reconfigure(encoding='utf-8')
sys.stderr.reconfigure(encoding='utf-8')

import os
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ─────────────────────────────────────────────
# 样式工具函数
# ─────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    """设置表格单元格背景色"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


FONT_NAME = '微软雅黑'


def _set_run_font(run):
    """统一设置 run 的字体为微软雅黑"""
    run.font.name = FONT_NAME
    run._element.rPr.get_or_add_rFonts().set(qn('w:eastAsia'), FONT_NAME)
    run._element.rPr.get_or_add_rFonts().set(qn('w:ascii'), FONT_NAME)
    run._element.rPr.get_or_add_rFonts().set(qn('w:hAnsi'), FONT_NAME)


def _clear_heading_theme_color(p):
    """清除标题段落中所有 run 的主题色，确保颜色设置生效"""
    for run in p.runs:
        rPr = run._element.get_or_add_rPr()
        # 移除主题色节点
        for color_el in rPr.findall(qn('w:color')):
            rPr.remove(color_el)


def add_heading(doc, text, level=1, color=None):
    """添加标题，强制字体微软雅黑，颜色默认黑色"""
    p = doc.add_heading(text, level=level)
    _clear_heading_theme_color(p)
    rgb = color if color else (0, 0, 0)
    for run in p.runs:
        run.font.color.rgb = RGBColor(*rgb)
        _set_run_font(run)
    return p


def add_paragraph(doc, text, bold=False, italic=False, size=11, color=None, indent=0):
    """添加普通段落"""
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)
    _set_run_font(run)
    return p


def add_placeholder(doc, text="【待填写】"):
    """添加占位符段落（橙色标注）"""
    p = doc.add_paragraph()
    run = p.add_run(f"📌 {text}")
    run.font.color.rgb = RGBColor(204, 102, 0)
    run.bold = True
    run.font.size = Pt(11)
    _set_run_font(run)
    return p


def add_analysis_text(doc, text):
    """添加AI分析文字段落"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run("🤖 分析：")
    run.bold = True
    run.font.size = Pt(10.5)
    run.font.color.rgb = RGBColor(70, 130, 180)
    _set_run_font(run)

    run2 = p.add_run(text)
    run2.font.size = Pt(10.5)
    run2.font.color.rgb = RGBColor(80, 80, 80)
    _set_run_font(run2)
    return p


def fmt_val(v, decimal=2, pct=False):
    """格式化数值为字符串"""
    if v is None:
        return "-"
    try:
        v = float(v)
    except (TypeError, ValueError):
        return str(v)
    if pct:
        return f"{v:.{decimal}f}%"
    if abs(v) >= 1e8:
        return f"{v/1e8:.{decimal}f}亿"
    if abs(v) >= 1e4:
        return f"{v/1e4:.{decimal}f}万"
    return f"{v:.{decimal}f}"


def calc_pct_change(cur, prev):
    """计算百分比变化"""
    try:
        if prev and prev != 0:
            return (cur - prev) / abs(prev) * 100
    except Exception:
        pass
    return None


def change_str(cur_val, prev_val):
    """返回环比字符串"""
    chg = calc_pct_change(cur_val, prev_val)
    if chg is None:
        return "-"
    arrow = "↑" if chg > 0 else "↓"
    return f"{arrow}{abs(chg):.1f}%"


# ─────────────────────────────────────────────
# 表格构建函数
# ─────────────────────────────────────────────

def make_table(doc, headers, rows, col_widths=None):
    """
    通用表格构建
    headers: list of str
    rows: list of list
    """
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 表头
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        r = hdr_cells[i].paragraphs[0].runs[0]
        r.bold = True
        r.font.color.rgb = RGBColor(255, 255, 255)
        r.font.size = Pt(10)
        _set_run_font(r)
        hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_bg(hdr_cells[i], "1F4E79")

    # 数据行
    for ri, row in enumerate(rows):
        cells = table.add_row().cells
        bg = "EBF3FB" if ri % 2 == 0 else "FFFFFF"
        for ci, val in enumerate(row):
            cells[ci].text = str(val) if val is not None else "-"
            cells[ci].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = cells[ci].paragraphs[0].runs[0]
            r.font.size = Pt(10)
            _set_run_font(r)
            set_cell_bg(cells[ci], bg)

    # 列宽
    if col_widths:
        for i, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(width)

    doc.add_paragraph()
    return table


# ─────────────────────────────────────────────
# 各章节构建函数
# ─────────────────────────────────────────────

def build_cover(doc, product_name, current_start, current_end):
    """封面"""
    doc.add_paragraph()
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(product_name.upper())
    run.font.size = Pt(28)
    run.bold = True
    run.font.color.rgb = RGBColor(31, 78, 121)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run(f"版本数据分析报告")
    run2.font.size = Pt(20)
    run2.font.color.rgb = RGBColor(89, 89, 89)

    doc.add_paragraph()

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run3 = p3.add_run(f"{current_start} 至 {current_end}")
    run3.font.size = Pt(14)
    run3.font.color.rgb = RGBColor(120, 120, 120)

    p4 = doc.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run4 = p4.add_run(f"生成日期：{datetime.now().strftime('%Y-%m-%d')}")
    run4.font.size = Pt(12)
    run4.font.color.rgb = RGBColor(150, 150, 150)

    doc.add_page_break()


def _period_days(meta, which="current"):
    """计算版本周期天数"""
    from datetime import date as _date
    try:
        if which == "current":
            s = meta.get("current_start")
            e = meta.get("current_end")
        else:
            s = meta.get("prev_start")
            e = meta.get("prev_end")
        if s and e:
            if not hasattr(s, 'toordinal'):
                s = _date.fromisoformat(str(s))
            if not hasattr(e, 'toordinal'):
                e = _date.fromisoformat(str(e))
            return (e - s).days + 1
    except Exception:
        pass
    return None


def build_chapter1_基础大盘(doc, all_data, analysis_results):
    """第一章：通用数据概览"""
    add_heading(doc, "第一章  通用数据概览", level=1, color=(0, 0, 0))

    meta = all_data.get("_meta", {})
    cur_days = _period_days(meta, "current")
    prev_days = _period_days(meta, "prev")
    periods_equal = (cur_days == prev_days)

    # 1.1 基础大盘
    add_heading(doc, "1.1 基础大盘数据", level=2)

    # 说明期间天数（如果不同则提示环比基于日均）
    if cur_days and prev_days:
        note = f"本版本 {cur_days} 天 | 上版本 {prev_days} 天"
        if not periods_equal:
            note += "（周期不同，收入/成本类环比基于日均）"
        add_paragraph(doc, note, italic=True, size=9, color=(120, 120, 120))

    base = all_data.get("基础数据", {})
    cur = base.get("current", {})
    prev = base.get("previous", {})

    # (显示标签, 显示用sub_key, 环比用sub_key, 是否百分比)
    # 对于总量类指标（收入/成本/DNU），显示 sum，环比用 mean（日均环比）
    metrics = [
        ("DAU", "mean", "mean", False),
        ("MAU", "mean", "mean", False),
        ("DNU", "sum", "mean", False),
        ("次留(%)", "mean", "mean", True),
        ("付费率(%)", "mean", "mean", True),
        ("破产率(%)", "mean", "mean", True),
        ("总收入($)", "sum", "mean", False),
        ("内购收入($)", "sum", "mean", False),
        ("广告收入($)", "sum", "mean", False),
        ("内购ARPU($)", "mean", "mean", False),
        ("内购ARPPU($)", "mean", "mean", False),
        ("推广成本($)", "sum", "mean", False),
    ]

    headers = ["指标", "本版本", "上版本", "环比"]
    rows = []
    for label, display_sub, cmp_sub, is_pct in metrics:
        c_val = cur.get(label, {}).get(display_sub) if cur else None
        p_val = prev.get(label, {}).get(display_sub) if prev else None
        c_cmp = cur.get(label, {}).get(cmp_sub) if cur else None
        p_cmp = prev.get(label, {}).get(cmp_sub) if prev else None
        rows.append([
            label,
            fmt_val(c_val, pct=is_pct),
            fmt_val(p_val, pct=is_pct),
            change_str(c_cmp, p_cmp) if (c_cmp is not None and p_cmp is not None) else "-"
        ])
    make_table(doc, headers, rows, col_widths=[5, 3.5, 3.5, 3])

    add_analysis_text(doc, analysis_results.get("基础大盘", ""))

    # 1.2 广告数据
    add_heading(doc, "1.2 广告数据", level=2)
    ad = all_data.get("广告大盘数据", {})
    ad_cur = ad.get("current", {})
    ad_prev = ad.get("previous", {})

    ad_metrics = [
        ("广告总收入", "sum", "mean", False),
        ("广告观看次数", "sum", "mean", False),
        ("广告观看人数", "sum", "mean", False),
        ("ECPM", "mean", "mean", False),
        ("广告arpu", "mean", "mean", False),
        ("广告渗透率", "mean", "mean", True),
        ("活跃人均次数", "mean", "mean", False),
    ]
    rows = []
    for label, display_sub, cmp_sub, is_pct in ad_metrics:
        c_val = ad_cur.get(label, {}).get(display_sub) if ad_cur else None
        p_val = ad_prev.get(label, {}).get(display_sub) if ad_prev else None
        c_cmp = ad_cur.get(label, {}).get(cmp_sub) if ad_cur else None
        p_cmp = ad_prev.get(label, {}).get(cmp_sub) if ad_prev else None
        rows.append([
            label,
            fmt_val(c_val, pct=is_pct),
            fmt_val(p_val, pct=is_pct),
            change_str(c_cmp, p_cmp) if (c_cmp is not None and p_cmp is not None) else "-"
        ])
    make_table(doc, headers, rows, col_widths=[5, 3.5, 3.5, 3])
    add_analysis_text(doc, analysis_results.get("广告数据", ""))

    # 1.3 活跃数据
    add_heading(doc, "1.3 活跃数据", level=2)
    act = all_data.get("活跃数据监控", {})
    act_cur = act.get("current", {})
    act_prev = act.get("previous", {})

    act_metrics = [
        ("spin用户数", "mean", False),
        ("人均spin次数", "mean", False),
        ("机台通过率", "mean", True),
        ("激励广告覆盖率", "mean", True),
        ("激励广告人均", "mean", False),
        ("插屏广告覆盖率", "mean", True),
        ("插屏广告人均", "mean", False),
    ]
    rows = []
    for label, sub_key, is_pct in act_metrics:
        c_val = act_cur.get(label, {}).get(sub_key) if act_cur else None
        p_val = act_prev.get(label, {}).get(sub_key) if act_prev else None
        rows.append([
            label,
            fmt_val(c_val, pct=is_pct),
            fmt_val(p_val, pct=is_pct),
            change_str(c_val, p_val) if (c_val and p_val) else "-"
        ])
    make_table(doc, headers, rows, col_widths=[5, 3.5, 3.5, 3])
    add_analysis_text(doc, analysis_results.get("活跃数据", ""))


def build_chapter2_新用户(doc, all_data, analysis_results):
    """第二章：新用户数据"""
    add_heading(doc, "第二章  新用户数据", level=1, color=(0, 0, 0))

    # 2.1 破冰率
    add_heading(doc, "2.1 注册破冰率漏斗", level=2)
    breakout = all_data.get("注册破冰率监控", {})
    stage = breakout.get("stage", {})

    if stage:
        day_keys = ["当日", "第1日", "第2日", "第3日", "第4日", "第5日", "第6日", "第7日"]
        headers = ["阶段值"] + [k for k in day_keys if k in stage]
        row = ["破冰率"] + [fmt_val(stage.get(k), pct=True) for k in day_keys if k in stage]
        make_table(doc, headers, [row])
    else:
        add_paragraph(doc, "（破冰率数据未找到阶段值）", italic=True)

    # 逐日明细
    cur_breakout = breakout.get("current", {})
    if cur_breakout and cur_breakout.get("daily_avg"):
        daily = cur_breakout["daily_avg"]
        add_paragraph(doc, "逐日均值：", bold=True)
        day_keys = [k for k in ["当日", "第1日", "第2日", "第3日", "第4日", "第5日", "第6日", "第7日"] if k in daily]
        row = [fmt_val(daily.get(k), pct=True) for k in day_keys]
        make_table(doc, day_keys, [row])

    add_analysis_text(doc, analysis_results.get("新用户", ""))

    # 2.2 首日人均Spin
    add_heading(doc, "2.2 首日人均 Spin（分渠道）", level=2)
    spin = all_data.get("首日人均spin", {})
    spin_cur = spin.get("current", {})

    if spin_cur:
        headers_spin = ["渠道", "人均Spin次数", "注册Spin比例", "秦始皇机台人均Spin"]
        rows = []
        for channel, metrics in spin_cur.items():
            rows.append([
                channel,
                fmt_val(metrics.get("spin人均次数", {}).get("mean")),
                fmt_val(metrics.get("注册spin比例", {}).get("mean"), pct=True),
                fmt_val(metrics.get("秦始皇机台人均spin次数", {}).get("mean")),
            ])
        make_table(doc, headers_spin, rows, col_widths=[4, 3.5, 3.5, 4])
    else:
        add_paragraph(doc, "（首日Spin数据缺失）", italic=True)


def build_chapter3_用户结构(doc, all_data, analysis_results):
    """第三章：用户结构"""
    add_heading(doc, "第三章  用户结构", level=1, color=(0, 0, 0))
    add_heading(doc, "3.1 生命周期 × 付费段分布", level=2)

    user = all_data.get("用户构成", {})
    cur = user.get("current", {})

    if cur:
        # 按付费段汇总
        r_totals = {}
        for key, val in cur.items():
            parts = key.split(" | ")
            if len(parts) == 2:
                r_seg = parts[1]
                r_totals[r_seg] = r_totals.get(r_seg, 0) + (val or 0)

        total = sum(r_totals.values()) or 1
        headers = ["付费段", "平均用户数", "占比"]
        rows = []
        for r, v in sorted(r_totals.items(), key=lambda x: x[1], reverse=True):
            rows.append([r, fmt_val(v), fmt_val(v / total * 100, pct=True)])
        make_table(doc, headers, rows, col_widths=[5, 4, 4])
    else:
        add_paragraph(doc, "（用户构成数据缺失）", italic=True)

    add_analysis_text(doc, analysis_results.get("用户结构", ""))


def build_chapter4_付费分析(doc, all_data, analysis_results):
    """第四章：付费分析"""
    add_heading(doc, "第四章  付费分析", level=1, color=(0, 0, 0))

    # 4.1 付费点位分布
    add_heading(doc, "4.1 付费点位分布（按收入排序）", level=2)
    pay_dist = all_data.get("付费点位分布", {})
    cur = pay_dist.get("current", {})

    if cur:
        headers = ["付费入口", "内购金额($)", "内购次数", "内购人数"]
        rows = []
        for pos, metrics in cur.items():
            rev = metrics.get("内购金额", 0)
            if rev == 0:
                continue  # 跳过零收入点位
            rows.append([
                pos,
                fmt_val(rev),
                fmt_val(metrics.get("内购次数", 0), decimal=0),
                fmt_val(metrics.get("内购人数", 0), decimal=0),
            ])
        make_table(doc, headers, rows, col_widths=[5.5, 3, 3, 3])
    else:
        add_paragraph(doc, "（付费点位数据缺失）", italic=True)

    add_analysis_text(doc, analysis_results.get("付费点位", ""))

    # 4.2 付费点位点击率（Top条目）
    add_heading(doc, "4.2 付费点位点击率（Top15）", level=2)
    ctr = all_data.get("付费点位点击率", {})
    ctr_cur = ctr.get("current", {})

    if ctr_cur:
        sorted_ctr = sorted(ctr_cur.items(), key=lambda x: x[1].get("avg_rate", 0), reverse=True)[:15]
        headers = ["入口 | 付费段", "平均点击率"]
        rows = [[k, fmt_val(v.get("avg_rate", 0), pct=True)] for k, v in sorted_ctr]
        make_table(doc, headers, rows, col_widths=[10, 3])
    else:
        add_paragraph(doc, "（点击率数据缺失）", italic=True)


def _r_label(r):
    """将任意格式的R段 key 转为可读标签"""
    r_str = str(r)
    if "500" in r_str and ("∞" in r_str or "inf" in r_str.lower() or "+" in r_str):
        return "大R(≥500)"
    if "100" in r_str and "500" in r_str:
        return "中R(100-500)"
    if "10" in r_str and "100" in r_str:
        return "小R(10-100)"
    if "0.1" in r_str and "10" in r_str:
        return "微R(<10)"
    return r_str


def _r_sort_key(r):
    """按大R→中R→小R→微R排序"""
    r_str = str(r)
    if "500" in r_str and ("∞" in r_str or "+" in r_str):
        return 0
    if "100" in r_str and "500" in r_str:
        return 1
    if "10" in r_str and "100" in r_str:
        return 2
    return 3


def build_chapter5_资源经济(doc, all_data, analysis_results):
    """第五章：资源经济健康"""
    add_heading(doc, "第五章  资源经济健康", level=1, color=(0, 0, 0))

    # 5.1 付费用户资源监控
    add_heading(doc, "5.1 付费用户资源监控（分R段）", level=2)
    res = all_data.get("付费用户资源监控", {})
    cur = res.get("current", {})

    if cur:
        headers = ["R段", "钻石消耗总和", "净消耗总额(无补充)", "净消耗中位数",
                   "钻石付费收入($)", "金币付费收入($)", "Spin赢得金币总量"]
        rows = []
        for r in sorted(cur.keys(), key=_r_sort_key):
            d = cur[r]
            rows.append([
                _r_label(r),
                fmt_val(d.get("消耗钻石总和")),
                fmt_val(d.get("钻石净消耗总额（无付费补充）")),
                fmt_val(d.get("钻石净消耗中位数（无付费补充）")),
                fmt_val(d.get("付费购买钻石购买价格总和")),
                fmt_val(d.get("付费购买金币购买价格总和")),
                fmt_val(d.get("下线事件.本次Spin赢得金币数量总和")),
            ])
        make_table(doc, headers, rows, col_widths=[3, 2.5, 3, 2.5, 2.5, 2.5, 3])
    else:
        add_paragraph(doc, "（资源监控数据缺失）", italic=True)

    add_analysis_text(doc, analysis_results.get("资源经济", ""))

    # 5.2 金币库存监控
    add_heading(doc, "5.2 付费用户金币库存（各R段中位数）", level=2)
    inv = all_data.get("付费用户金币库存监控", {})
    inv_cur = inv.get("current", {})

    if inv_cur:
        headers = ["R段", "当前资产中位数（金币）", "Spin赢得金币中位数"]
        rows = []
        for r in sorted(inv_cur.keys(), key=_r_sort_key):
            d = inv_cur[r]
            rows.append([
                _r_label(r),
                fmt_val(d.get("用户登录.当前资产中位数")),
                fmt_val(d.get("下线事件.本次Spin赢得金币数量中位数")),
            ])
        make_table(doc, headers, rows, col_widths=[4, 5, 5])
    else:
        add_paragraph(doc, "（金币库存数据缺失）", italic=True)


def build_chapter6_版本内容(doc):
    """第六章：版本更新内容（占位）"""
    add_heading(doc, "第六章  版本更新内容复盘", level=1, color=(0, 0, 0))
    add_paragraph(doc, "本章节需根据本版本具体更新内容手动填写。以下为参考结构：", italic=True, color=(120, 120, 120))
    doc.add_paragraph()

    for i in range(1, 4):
        add_heading(doc, f"6.{i} 更新内容 {i}", level=2)
        add_placeholder(doc, f"请填写更新内容{i}名称及背景说明")
        add_placeholder(doc, "关键数据指标（参考相关报表数据）")
        add_placeholder(doc, "结论与建议")
        doc.add_paragraph()


def build_chapter7_综合总结(doc, analysis_results):
    """第七章：综合总结"""
    add_heading(doc, "第七章  综合总结", level=1, color=(0, 0, 0))

    summary = analysis_results.get("综合总结", "")
    if summary:
        # 按行拆分输出
        lines = [l.strip() for l in summary.split("\n") if l.strip()]
        for line in lines:
            p = doc.add_paragraph(style='List Bullet')
            run = p.add_run(line)
            run.font.size = Pt(11)
    else:
        add_placeholder(doc, "综合总结生成失败，请手动填写")


# ─────────────────────────────────────────────
# 主入口
# ─────────────────────────────────────────────

def build_report(all_data, analysis_results, product_name, output_dir):
    """
    生成完整Word报告
    返回输出文件路径
    """
    os.makedirs(output_dir, exist_ok=True)

    meta = all_data.get("_meta", {})
    current_start = meta.get("current_start")
    current_end = meta.get("current_end")

    doc = Document()

    # 设置页面边距
    from docx.oxml import OxmlElement
    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # 全局字体统一为微软雅黑（样式级别兜底）
    for style_name in ['Normal', 'Heading 1', 'Heading 2', 'Heading 3',
                       'Heading 4', 'List Bullet', 'List Bullet 2']:
        try:
            style = doc.styles[style_name]
            style.font.name = FONT_NAME
            rPr = style._element.get_or_add_rPr()
            rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)
            rPr.rFonts.set(qn('w:ascii'), FONT_NAME)
            rPr.rFonts.set(qn('w:hAnsi'), FONT_NAME)
        except KeyError:
            pass

    print("\n[START] 开始生成Word报告...")

    build_cover(doc, product_name, current_start, current_end)
    print("  [OK] 封面")

    build_chapter1_基础大盘(doc, all_data, analysis_results)
    print("  [OK] 第一章：通用数据概览")

    doc.add_page_break()
    build_chapter2_新用户(doc, all_data, analysis_results)
    print("  [OK] 第二章：新用户数据")

    doc.add_page_break()
    build_chapter3_用户结构(doc, all_data, analysis_results)
    print("  [OK] 第三章：用户结构")

    doc.add_page_break()
    build_chapter4_付费分析(doc, all_data, analysis_results)
    print("  [OK] 第四章：付费分析")

    doc.add_page_break()
    build_chapter5_资源经济(doc, all_data, analysis_results)
    print("  [OK] 第五章：资源经济健康")

    doc.add_page_break()
    build_chapter6_版本内容(doc)
    print("  [OK] 第六章：版本内容占位")

    doc.add_page_break()
    build_chapter7_综合总结(doc, analysis_results)
    print("  [OK] 第七章：综合总结")

    # 输出文件名
    start_str = str(current_start).replace("-", "")
    end_str = str(current_end).replace("-", "")
    safe_name = product_name.replace(" ", "_").upper()
    filename = f"{safe_name}_{start_str}-{end_str}_数据分析报告.docx"
    output_path = os.path.join(output_dir, filename)

    doc.save(output_path)
    print(f"\n[SAVED] 报告已保存：{output_path}")
    return output_path


# ─────────────────────────────────────────────
# 命令行入口（供 Claude 直接调用）
# ─────────────────────────────────────────────

def _load_json_data(data_path, analysis_path):
    """
    从 JSON 文件重建 all_data 和 analysis_results
    data_path: data_extractor.py 输出的 extracted_data.json
    analysis_path: Claude 写入的 analysis_results.json
    """
    import json
    from datetime import date

    with open(data_path, encoding="utf-8") as f:
        extracted = json.load(f)

    with open(analysis_path, encoding="utf-8") as f:
        analysis_results = json.load(f)

    meta = extracted.get("meta", {})

    def parse_date(s):
        if not s or s == "None":
            return None
        return date.fromisoformat(s)

    # 重建 all_data，只需要 _meta 和各章节聚合值供表格渲染
    # report_builder 的表格构建函数从 all_data 读数据
    # 我们把 extracted["raw"] 中的数据还原回 all_data 格式
    raw = extracted.get("raw", {})

    all_data = {
        "_meta": {
            "current_start": parse_date(meta.get("current_start")),
            "current_end": parse_date(meta.get("current_end")),
            "prev_start": parse_date(meta.get("prev_start")),
            "prev_end": parse_date(meta.get("prev_end")),
        }
    }

    # 从 sections 中重建基础数据结构（供表格用）
    # 直接把 raw 中的数据透传
    for k, v in raw.items():
        all_data[k] = v

    # 基础数据需要从 sections 文本中补充（表格从 all_data 读）
    # 这里提供一个从 extracted_data.json 恢复基础数据的方式
    # 由于基础数据原始结构太大没有序列化，用 sections 文本里的数据
    # 表格渲染时会尝试读取 all_data["基础数据"]，如果不存在则跳过
    all_data["基础数据"] = raw.get("基础数据", {})

    return all_data, analysis_results


if __name__ == "__main__":
    import argparse
    import json

    parser = argparse.ArgumentParser(description="版本复盘报告生成器")
    parser.add_argument("--data", required=True, help="data_extractor 输出的 JSON 路径")
    parser.add_argument("--analysis", required=True, help="Claude 写入的分析结果 JSON 路径")
    parser.add_argument("--product", required=True, help="产品名称")
    parser.add_argument("--start", required=True, help="版本开始日期 YYYY-MM-DD")
    parser.add_argument("--end", required=True, help="版本结束日期 YYYY-MM-DD")
    parser.add_argument("--output", required=True, help="报告输出目录")
    args = parser.parse_args()

    print(f"[START] 加载数据文件：{args.data}")
    all_data, analysis_results = _load_json_data(args.data, args.analysis)

    # 用命令行参数覆盖 meta 日期（确保准确）
    from datetime import date as _date
    all_data["_meta"]["current_start"] = _date.fromisoformat(args.start)
    all_data["_meta"]["current_end"] = _date.fromisoformat(args.end)

    output_path = build_report(all_data, analysis_results, args.product, args.output)
    print(f"[DONE] 报告路径：{output_path}")
